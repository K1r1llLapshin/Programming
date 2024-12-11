from docx import Document, shared

document = Document('./Lab10/ATmega.docx')

pr = document.add_paragraph()
pr.alignment = 1
pr.add_run().add_picture('./Lab10/ATmega328.png', width= shared.Cm(5), height= shared.Cm(5))
document.add_paragraph('ATmega328').alignment = 1


document.save('./Lab10/ATmega.docx')