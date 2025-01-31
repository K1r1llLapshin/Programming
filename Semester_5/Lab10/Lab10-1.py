from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

document = Document()

document.add_paragraph('В микроконтроллерах ATmega, используемых на платформах Arduino, существует три вида памяти:')

document.add_paragraph('Флеш-пямять: используется для хранения скетчей', style='List Bullet 3')
l2 = document.add_paragraph(style='List Bullet 3')
l2.add_run('ОЗУ (')
l2.add_run('SRAM').bold = True
l2.add_run(' — ')
l2.add_run('static random access memory').italic = True
l2.add_run(', статическая опреативная память с произвольным доступом): используется для храниения и работы переменных.')
document.add_paragraph('EEPROM (энергонезависемая память): используется для зранения постоянной информации.', style='List Bullet 3')

document.add_paragraph('Флеш-память и EEPROM являются энергонезависимыми видами памяти (данные сохраняются при отключении питания). ОЗУ является энергозависимой памятью.')

table = document.add_table(cols=5, rows=4)
table.style = 'Table Grid'

table_inf = [['', 'ATmega168', 'ATmega328', 'ATmega1280', 'ATmega2560'],
         ['Flash (1 кБ flash-памяти занят загрузчиком)', '16 Кбайт','32 Кбайт','128 Кбайт', '256 Кбайт'],
         ['SRAM', '1 Кбайт', '2 Кбайт', '8 Кбайт', '8 Кбайт'],
         ['EEPROM', '512 байт', '1024 байта', '4 Кбайт', '4 Кбайт']]
for col in range (5):
    for row in range(4):
        cell = table.cell(row, col)
        if col < 1 or row < 1:
            cell.paragraphs[0].add_run(table_inf[row][col]).bold =True
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            cell.text = table_inf[row][col]

document.add_paragraph()
run = '''Память EEPROM, по заявлениям производителя, обладает гарантированным жизненным циклом 100 000 операций записи/стирания и 100 лет хранения данных при температуре 25 С. Эти данные не распространяются на операции чтения данных из EEPROM — чтение данных не лимитировано. Исходя из этого, нужно проектировать свои скетчи максимально щадящими по отношению к EEPROM.'''
document.add_paragraph().add_run(run).italic = True

document.save('./Lab10/ATmega.docx')